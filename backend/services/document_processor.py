"""Document processing service for PDF and DOCX files."""
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
from typing import List, Dict, Tuple
import spacy
from pathlib import Path
import logging
from ..config import get_settings

logger = logging.getLogger(__name__)
settings = get_settings()


class DocumentProcessor:
    """Process PDF and DOCX documents."""
    
    def __init__(self):
        """Initialize document processor with spaCy."""
        try:
            self.nlp = spacy.load(settings.spacy_model)
        except OSError:
            logger.warning(f"spaCy model {settings.spacy_model} not found. Downloading...")
            import subprocess
            subprocess.run(["python", "-m", "spacy", "download", settings.spacy_model])
            self.nlp = spacy.load(settings.spacy_model)
            
        # Disable unnecessary pipeline components for speed
        self.nlp.disable_pipes([pipe for pipe in self.nlp.pipe_names if pipe not in ["tok2vec", "tagger"]])
        
    def extract_text_from_pdf(self, file_path: str) -> Tuple[str, Dict]:
        """
        Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Tuple of (full_text, metadata)
        """
        full_text = []
        metadata = {"pages": 0, "page_texts": []}
        
        try:
            # Use pdfplumber for better text extraction
            with pdfplumber.open(file_path) as pdf:
                metadata["pages"] = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text()
                    if text:
                        full_text.append(text)
                        metadata["page_texts"].append({
                            "page": page_num,
                            "text": text,
                            "char_count": len(text)
                        })
                        
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata["pages"] = len(pdf_reader.pages)
                    
                    for page_num, page in enumerate(pdf_reader.pages, start=1):
                        text = page.extract_text()
                        if text:
                            full_text.append(text)
                            metadata["page_texts"].append({
                                "page": page_num,
                                "text": text,
                                "char_count": len(text)
                            })
            except Exception as e2:
                logger.error(f"Fallback PDF extraction also failed: {e2}")
                raise
        
        return "\n\n".join(full_text), metadata
    
    def extract_text_from_docx(self, file_path: str) -> Tuple[str, Dict]:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Tuple of (full_text, metadata)
        """
        try:
            doc = DocxDocument(file_path)
            
            paragraphs = []
            metadata = {
                "paragraphs": 0,
                "tables": len(doc.tables),
                "sections": len(doc.sections)
            }
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            metadata["paragraphs"] = len(paragraphs)
            
            # Extract tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    if row_text.strip():
                        table_texts.append(row_text)
            
            # Combine paragraphs and tables
            full_text = "\n\n".join(paragraphs)
            if table_texts:
                full_text += "\n\n" + "\n".join(table_texts)
            
            return full_text, metadata
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise
    
    def chunk_text(self, text: str, page_number: int = None) -> List[Dict]:
        """
        Chunk text using token-based approach with spaCy.
        
        Args:
            text: Text to chunk
            page_number: Optional page number for PDF chunks
            
        Returns:
            List of chunk dictionaries with content and metadata
        """
        if not text.strip():
            return []
        
        # Process text with spaCy
        doc = self.nlp(text)
        
        chunks = []
        current_chunk = []
        current_token_count = 0
        
        for sent in doc.sents:
            sent_tokens = len([token for token in sent if not token.is_space])
            
            # If adding this sentence exceeds chunk size, save current chunk
            if current_token_count + sent_tokens > settings.chunk_size and current_chunk:
                chunk_text = " ".join(current_chunk)
                chunks.append({
                    "content": chunk_text,
                    "token_count": current_token_count,
                    "page_number": page_number,
                    "metadata": {}
                })
                
                # Start new chunk with overlap
                if settings.chunk_overlap > 0:
                    # Keep last few sentences for overlap
                    overlap_tokens = 0
                    overlap_sents = []
                    for sent_text in reversed(current_chunk):
                        sent_token_count = len(self.nlp(sent_text))
                        if overlap_tokens + sent_token_count <= settings.chunk_overlap:
                            overlap_sents.insert(0, sent_text)
                            overlap_tokens += sent_token_count
                        else:
                            break
                    current_chunk = overlap_sents
                    current_token_count = overlap_tokens
                else:
                    current_chunk = []
                    current_token_count = 0
            
            # Add sentence to current chunk
            current_chunk.append(sent.text)
            current_token_count += sent_tokens
        
        # Add final chunk
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            chunks.append({
                "content": chunk_text,
                "token_count": current_token_count,
                "page_number": page_number,
                "metadata": {}
            })
        
        return chunks
    
    def process_document(self, file_path: str, file_type: str) -> Tuple[List[Dict], Dict]:
        """
        Process document and return chunks with metadata.
        
        Args:
            file_path: Path to document
            file_type: "pdf" or "docx"
            
        Returns:
            Tuple of (chunks, document_metadata)
        """
        # Extract text
        if file_type.lower() == "pdf":
            full_text, doc_metadata = self.extract_text_from_pdf(file_path)
            
            # Chunk by page for PDFs
            all_chunks = []
            for page_info in doc_metadata.get("page_texts", []):
                page_chunks = self.chunk_text(page_info["text"], page_info["page"])
                all_chunks.extend(page_chunks)
                
        elif file_type.lower() in ["docx", "doc"]:
            full_text, doc_metadata = self.extract_text_from_docx(file_path)
            all_chunks = self.chunk_text(full_text)
            
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        logger.info(f"Processed {file_type} document: {len(all_chunks)} chunks created")
        
        return all_chunks, doc_metadata


# Global instance
_document_processor = None


def get_document_processor() -> DocumentProcessor:
    """Get singleton document processor instance."""
    global _document_processor
    if _document_processor is None:
        _document_processor = DocumentProcessor()
    return _document_processor
